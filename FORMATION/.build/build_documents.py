from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\PROJETS\madgi-esr")
OUT = ROOT / "FORMATION"
OUT.mkdir(parents=True, exist_ok=True)
NAVY = "173A6A"; BLUE = "2E74B5"; GREEN = "147A4B"; GOLD = "D89B2B"; PALE = "E8EEF5"; LIGHT = "F5F7FA"; RED = "B42318"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'),'dxa')

def setup(doc, title, subtitle, code):
    sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=Inches(0.75); sec.left_margin=sec.right_margin=Inches(0.85)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string('26354A')
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
    for name,size,color,before,after in [('Title',25,NAVY,0,12),('Heading 1',16,BLUE,16,8),('Heading 2',13,GREEN,12,6),('Heading 3',11,GOLD,9,4)]:
        s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    header=sec.header.paragraphs[0]; header.text=f"MADGI • ÉPARGNE SANTÉ RETRAITE   |   {code}"; header.style=styles['Caption']; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs: run.font.color.rgb=RGBColor.from_string(BLUE); run.font.bold=True
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Support de formation — Usage interne • ")
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    logo=ROOT/'src/assets/logos/logo-madgi.jpg'
    if logo.exists(): p.add_run().add_picture(str(logo), width=Inches(1.25))
    p=doc.add_paragraph(title, style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph(subtitle); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string(GREEN); r.font.bold=True
    p=doc.add_paragraph('Version formation • 19 août 2026'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.color.rgb=RGBColor.from_string('667085')
    doc.add_paragraph()

def add_callout(doc, title, text, color=PALE):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(6.75)
    c=t.cell(0,0); shade(c,color); set_cell_margins(c,140,180,140,180)
    p=c.paragraphs[0]; r=p.add_run(title+' — '); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); p.add_run(text)
    doc.add_paragraph()

def bullets(doc, items, level=0):
    for item in items:
        p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.paragraph_format.space_after=Pt(3); p.add_run(item)

def numbered(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Number'); p.paragraph_format.space_after=Pt(3); p.add_run(item)

def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,PALE); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c)
        for r in c.paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph(); return t

def section(doc, heading, intro=None):
    doc.add_heading(heading, level=1)
    if intro: doc.add_paragraph(intro)

def save(doc, name):
    path=OUT/name; doc.save(path); return path

def program():
    d=Document(); setup(d,'PROGRAMME DE FORMATION','Présentation et prise en main opérationnelle de MADGI ESR','01')
    add_callout(d,'Finalité','Rendre les gestionnaires et administrateurs autonomes sur les opérations courantes, les contrôles et la clôture d’une période ESR.')
    section(d,'Public, durée et prérequis')
    table(d,['Élément','Cadre proposé'],[
        ['Public','Gestionnaires ESR, comptables, administrateurs, agents d’accueil et référents métier'],
        ['Durée','1 journée — 08 h 30 à 16 h 30, pauses comprises'],
        ['Format','30 % présentation, 30 % démonstration, 40 % pratique guidée'],
        ['Prérequis','Compte utilisateur actif, navigateur récent, accès à l’environnement de formation'],
        ['Matériel','Vidéoprojecteur, connexion réseau, un poste pour deux participants, jeux de données de test']], [1.45,5.2])
    section(d,'Objectifs pédagogiques')
    bullets(d,['Se connecter et identifier les fonctions autorisées selon son rôle.','Créer, rechercher et mettre à jour un adhérent sans doublon.','Traiter les cotisations trimestrielles, les régularisations et les versements spontanés.','Suivre correctement un paiement par chèque jusqu’à son encaissement.','Contrôler un compte individuel ESR, produire un reçu et exploiter les états.','Préparer puis clôturer une période en respectant les contrôles bloquants.'])
    section(d,'Déroulé détaillé')
    table(d,['Heure','Séquence','Résultat attendu'],[
        ['08:30–09:00','Accueil, objectifs, rôles et règles de sécurité','Comprendre le périmètre et les responsabilités'],
        ['09:00–09:45','Navigation, tableau de bord et recherche','Se repérer et retrouver une information'],
        ['09:45–10:30','Gestion des adhérents','Créer et contrôler une fiche'],
        ['10:45–12:00','Cotisations trimestrielles et précomptes','Générer, importer et régulariser'],
        ['13:00–14:00','Versement spontané et recalcul ESR','Saisir et expliquer l’impact financier'],
        ['14:00–14:45','Paiements directs et workflow des chèques','Faire progresser un chèque jusqu’à ENCAISSÉ'],
        ['15:00–15:45','Comptes, reçus, reporting et clôture','Contrôler et produire les justificatifs'],
        ['15:45–16:20','Exercices pratiques et QCM','Valider l’autonomie'],
        ['16:20–16:30','Synthèse, assistance et engagements','Fixer les actions post-formation']], [1.0,2.35,3.3])
    section(d,'Critères de réussite')
    bullets(d,['Présence à l’ensemble des séquences.','Score d’au moins 70 % au QCM.','Réalisation correcte d’au moins 3 scénarios pratiques sur 4.','Aucune validation fictive ou manipulation sur la production pendant les exercices.'])
    return save(d,'01_Programme_formation_MADGI_ESR.docx')

def demo():
    d=Document(); setup(d,'SCÉNARIO DE DÉMONSTRATION','Conducteur pas à pas pour une démonstration fluide et vérifiable','03')
    add_callout(d,'Règle de démonstration','Utiliser uniquement des adhérents de test. Annoncer l’objectif avant chaque manipulation et montrer le résultat attendu après chaque validation.')
    section(d,'Préparation des données')
    bullets(d,['Un compte GESTIONNAIRE et un compte ADMINISTRATEUR actifs.','Une période ESR ouverte avec au moins trois adhérents de test.','Un adhérent sans cotisation spontanée et un autre avec historique.','Un chèque de test, une référence de bordereau et une référence d’avis de crédit fictives.'])
    section(d,'Parcours narratif — 45 minutes')
    table(d,['Min.','Action à l’écran','Message à faire passer','Contrôle visible'],[
      ['0–5','Connexion et tableau de bord','Les menus et actions dépendent du rôle.','Nom, rôle et période affichés'],
      ['5–10','Recherche puis ouverture d’un adhérent','Toujours vérifier matricule et identité avant une opération.','Fiche et historique cohérents'],
      ['10–17','Afficher les précomptes du trimestre','La cotisation trimestrielle appartient à une période ouverte.','Période, montant, statut'],
      ['17–25','Saisir un versement spontané en espèces','Le versement augmente le capital et entraîne un recalcul; il ne soustrait pas mécaniquement son montant de l’échéance.','Montant > 0, période, date de valeur'],
      ['25–36','Saisir puis traiter un chèque','Une saisie n’est pas un encaissement. Le compte n’est crédité qu’à ENCAISSÉ.','Badges SAISI, CONTRÔLE, DÉPOSÉ, COMPENSÉ, VALIDÉ, ENCAISSÉ'],
      ['36–40','Télécharger le reçu','Le reçu est disponible après encaissement.','Dates au format JJ/MM/AAAA'],
      ['40–45','Contrôle de clôture','Les alertes doivent être levées avant clôture définitive.','Autorisation et liste des blocages']], [0.55,1.8,2.8,1.5])
    section(d,'Script exact du chèque')
    numbered(d,['Créer le paiement avec le moyen CHÈQUE et vérifier qu’il apparaît avec le statut SAISI.','Cliquer sur « Contrôler » après rapprochement avec les informations reçues.','Cliquer sur « Déposer en banque » et renseigner la référence du bordereau ainsi que la date de dépôt.','Après retour bancaire, cliquer sur « Confirmer compensation » et saisir la référence de l’avis de crédit et sa date.','Cliquer sur « Valider », puis sur « Encaisser ».','Vérifier le compte ESR, la cotisation recalculée et télécharger le reçu PDF.'])
    add_callout(d,'Cas d’impayé','Depuis le statut DÉPOSÉ EN BANQUE, utiliser « Impayé » et saisir le motif du rejet. Ne jamais forcer un chèque rejeté vers ENCAISSÉ.','FCE8E6')
    section(d,'Questions à poser au groupe')
    bullets(d,['À quel moment le capital acquis change-t-il ?','Quelles références doivent être conservées pour un chèque ?','Qui peut clôturer définitivement une période ?','Que vérifiez-vous avant de sélectionner un adhérent ?'])
    return save(d,'03_Scenario_demonstration_MADGI_ESR.docx')

def guide():
    d=Document(); setup(d,'GUIDE UTILISATEUR','Les opérations essentielles de MADGI ESR','04')
    add_callout(d,'Principe de sécurité','Chaque opération doit être effectuée avec un compte nominatif. Ne partagez jamais votre mot de passe et vérifiez l’adhérent, la période, le montant et le statut avant de valider.')
    for h,paras,items in [
      ('1. Se connecter et se repérer',['Après connexion, le tableau de bord présente les indicateurs utiles et le menu donne accès aux fonctions autorisées.'],['Vérifier votre nom et votre rôle.','Identifier la période ouverte avant toute opération de cotisation.','Utiliser la recherche par matricule dès que possible.']),
      ('2. Gérer les adhérents',['La fiche adhérent centralise l’identité, les paramètres de cotisation et l’historique.'],['Rechercher avant de créer afin d’éviter les doublons.','Contrôler matricule, nom, prénoms et situation administrative.','Enregistrer puis relire les informations sensibles.']),
      ('3. Traiter les cotisations trimestrielles',['Le module Précomptes permet de travailler sur une période ouverte, de générer les montants, d’intégrer les retours et de traiter les écarts.'],['Sélectionner la bonne période.','Ne jamais confondre montant généré, précompté, encaissé et régularisé.','Vérifier les lignes à montant nul avant la clôture.']),
      ('4. Enregistrer une cotisation spontanée',['Choisir l’adhérent, saisir un montant strictement positif, le moyen de paiement et la date. La date de valeur correspond au premier jour du trimestre suivant.'],['Le versement augmente le capital acquis.','La cotisation trimestrielle est recalculée à partir du capital restant à constituer et des trimestres restants.','Le montant spontané n’est pas simplement déduit de la cotisation trimestrielle courante.']),
      ('5. Confirmer un paiement par chèque',['Le chèque reste visible dans la liste des paiements avec son statut. Il doit franchir toutes les étapes avant d’être comptabilisé comme encaissé.'],['SAISI : paiement enregistré.','CONTRÔLÉ : informations vérifiées.','DÉPOSÉ EN BANQUE : bordereau et date renseignés.','COMPENSÉ : avis de crédit bancaire reçu.','VALIDÉ puis ENCAISSÉ : capitalisation effective et reçu disponible.']),
      ('6. Consulter le compte ESR et les reçus',['Le compte individuel permet de vérifier le capital acquis et les mouvements. Le reçu PDF constitue le justificatif utilisateur.'],['Comparer le mouvement au paiement encaissé.','Vérifier les dates affichées au format JJ/MM/AAAA.','Conserver le reçu selon les règles d’archivage internes.']),
      ('7. Clôturer une période',['La clôture est une opération définitive réservée aux profils autorisés. Le système affiche les contrôles et les alertes.'],['Traiter les paiements en attente et les anomalies.','Contrôler les précomptes, régularisations et montants nuls.','Ne clôturer que lorsque l’autorisation est affichée.'])]:
        section(d,h,paras[0]); bullets(d,items)
    section(d,'8. Messages d’erreur : premiers réflexes')
    table(d,['Situation','Action recommandée'],[
      ['Demande impossible à traiter','Noter l’heure, l’adhérent et l’action; actualiser une seule fois; transmettre la capture au support si l’erreur persiste.'],
      ['Paiement introuvable','Ouvrir la liste des paiements et vérifier les filtres et le statut, notamment SAISI ou DÉPOSÉ EN BANQUE.'],
      ['Bouton non disponible','Vérifier le rôle utilisateur et l’étape précédente du workflow.'],
      ['Clôture impossible','Lire chaque alerte, corriger les blocages puis relancer le contrôle.'],
      ['Montant inattendu','Contrôler période, capital acquis, paramètres de l’adhérent et nombre de trimestres restants.']], [2.0,4.65])
    return save(d,'04_Guide_utilisateur_MADGI_ESR.docx')

def fiches():
    d=Document(); setup(d,'FICHES PRATIQUES','Aide-mémoire à conserver au poste de travail','05')
    cards=[
      ('FICHE 1 — Cotisation spontanée',['Rechercher l’adhérent par matricule.','Vérifier la cotisation trimestrielle affichée.','Saisir un montant supérieur à zéro.','Choisir le moyen de versement et la date.','Valider puis vérifier la période et la date de valeur.','Contrôler le compte ESR et le nouveau montant trimestriel.']),
      ('FICHE 2 — Chèque',['Créer le paiement avec CHÈQUE.','Contrôler les informations.','Déposer en banque : saisir bordereau + date.','Confirmer la compensation : saisir avis de crédit + date.','Valider puis encaisser.','Télécharger le reçu; en cas de rejet, utiliser IMPAYÉ.']),
      ('FICHE 3 — Précompte trimestriel',['Vérifier la période ouverte.','Générer ou consulter les lignes.','Exporter selon la procédure interne.','Importer ou saisir les retours.','Régulariser les non-précomptés.','Contrôler les statuts et les montants avant clôture.']),
      ('FICHE 4 — Clôture',['Utiliser le contrôle de clôture.','Traiter chaque alerte affichée.','Vérifier qu’aucun paiement ne reste en attente.','Contrôler les montants nuls ou incohérents.','Faire valider les états de contrôle.','Clôturer avec un compte administrateur.'])]
    for idx,(title,steps) in enumerate(cards):
        section(d,title); numbered(d,steps)
        if idx<3: add_callout(d,'À ne pas oublier',['Le versement spontané déclenche un recalcul, pas une soustraction directe.','Un chèque n’est capitalisé qu’au statut ENCAISSÉ.','Une période clôturée ne doit plus recevoir d’opération.','La clôture est définitive et réservée aux profils habilités.'][idx])
    section(d,'Escalade au support')
    bullets(d,['Indiquer votre nom et votre rôle.','Donner le matricule de l’adhérent sans transmettre de données inutiles.','Préciser la période, le montant, l’heure et l’action réalisée.','Joindre une capture du message complet.','Ne pas multiplier les validations en cas de lenteur.'])
    return save(d,'05_Fiches_pratiques_MADGI_ESR.docx')

def evaluation():
    d=Document(); setup(d,'EXERCICES ET ÉVALUATION','Mise en situation, QCM et corrigé formateur','06')
    section(d,'Consignes')
    bullets(d,['Travaillez exclusivement dans l’environnement de formation.','Notez les contrôles effectués, pas seulement les clics.','Barème : QCM 10 points + pratique 20 points. Réussite à partir de 21/30.'])
    section(d,'Exercices pratiques')
    table(d,['Exercice','Mission','Preuves attendues','Pts'],[
      ['1','Retrouver un adhérent par matricule et vérifier sa période et sa cotisation.','Identité contrôlée; aucune modification inutile.','4'],
      ['2','Enregistrer une cotisation spontanée de 500 000 FCFA en espèces.','Montant positif; période/date de valeur correctes; nouveau calcul constaté.','6'],
      ['3','Traiter un chèque de SAISI à ENCAISSÉ.','Bordereau, avis de crédit, étapes respectées, reçu téléchargé.','6'],
      ['4','Analyser un contrôle de clôture et proposer les corrections.','Blocages identifiés; ordre d’action cohérent; pas de clôture forcée.','4']], [0.65,2.4,2.9,0.45])
    section(d,'QCM — une seule bonne réponse')
    qs=[
      '1. À quel statut un chèque alimente-t-il le capital ESR ?  A SAISI  B DÉPOSÉ  C COMPENSÉ  D ENCAISSÉ',
      '2. La cotisation spontanée :  A annule le trimestre  B est soustraite directement  C déclenche un recalcul  D ferme le compte',
      '3. Avant une saisie, le meilleur identifiant de recherche est :  A le prénom  B le matricule  C le montant  D la date',
      '4. Pour déposer un chèque, il faut renseigner :  A un mot de passe  B bordereau et date  C un reçu  D l’âge retraite',
      '5. Après compensation, l’étape suivante est :  A supprimer  B clôturer  C valider  D recréer',
      '6. Qui réalise la clôture définitive ?  A tout utilisateur  B l’adhérent  C un profil autorisé/administrateur  D la banque',
      '7. Une date de reçu doit être présentée sous la forme :  A JJ/MM/AAAA  B montant  C trimestre seul  D heure seule',
      '8. Si un chèque est rejeté par la banque :  A encaisser  B utiliser Impayé et saisir le motif  C supprimer  D ignorer',
      '9. Avant clôture, il faut :  A ignorer les alertes  B traiter les blocages  C changer d’adhérent  D imprimer seulement',
      '10. En cas d’erreur persistante :  A cliquer en boucle  B partager son mot de passe  C documenter puis contacter le support  D clôturer']
    for q in qs: d.add_paragraph(q)
    section(d,'Corrigé formateur')
    table(d,['Question','Réponse','Justification'],[[i+1,a,j] for i,(a,j) in enumerate([
      ('D','Capitalisation au statut ENCAISSÉ.'),('C','Recalcul actuariel à partir du capital restant.'),('B','Le matricule réduit le risque d’homonymie.'),('B','Traçabilité du dépôt bancaire.'),('C','COMPENSÉ précède VALIDÉ.'),('C','Action réservée aux profils habilités.'),('A','Format français demandé.'),('B','Le rejet bancaire doit rester traçable.'),('B','La clôture dépend des contrôles.'),('C','Éviter les doubles opérations et fournir les éléments utiles.')])], [0.8,1.0,4.2])
    return save(d,'06_Exercices_evaluation_MADGI_ESR.docx')

def checklist():
    d=Document(); setup(d,'CHECKLIST AVANT FORMATION','Contrôle technique, fonctionnel et logistique','07')
    add_callout(d,'Décision GO / NO GO','Ne pas commencer la démonstration tant qu’un compte de test, une période ouverte ou le workflow de paiement ne sont pas disponibles.')
    groups={
      'J-5 — Organisation':['Valider la liste des participants et leurs rôles.','Réserver salle, vidéoprojecteur et accès réseau.','Envoyer convocation, horaires et prérequis.','Préparer les supports imprimés ou numériques.'],
      'J-2 — Comptes et données':['Créer les comptes de formation.','Vérifier les droits GESTIONNAIRE et ADMINISTRATEUR.','Préparer les adhérents fictifs et une période ouverte.','Préparer un chèque fictif, bordereau et avis de crédit.'],
      'J-1 — Recette applicative':['Tester connexion/déconnexion.','Tester recherche et ouverture d’une fiche adhérent.','Tester une cotisation spontanée > 0.','Vérifier le recalcul de la cotisation trimestrielle.','Tester toutes les étapes du chèque jusqu’au reçu PDF.','Vérifier les dates JJ/MM/AAAA sur le reçu.','Tester le contrôle de clôture sans clôturer la période de démonstration.'],
      'Jour J — Salle':['Ouvrir l’application et le diaporama.','Vérifier zoom, résolution et lisibilité.','Couper notifications et données personnelles visibles.','Prévoir une copie locale des supports.','Afficher les objectifs et règles de sécurité.'],
      'Après session':['Collecter présence, scores et satisfaction.','Recenser les questions non résolues.','Transmettre guide et fiches pratiques.','Planifier un point de suivi à J+7.']}
    for h,items in groups.items():
        section(d,h)
        for x in items: d.add_paragraph('☐ '+x)
    section(d,'Feuille de décision')
    table(d,['Contrôle','État','Observation'],[['Application accessible','☐ GO  ☐ NO GO',''],['Comptes de test actifs','☐ GO  ☐ NO GO',''],['Données de démonstration prêtes','☐ GO  ☐ NO GO',''],['Workflow chèque validé','☐ GO  ☐ NO GO',''],['Supports disponibles','☐ GO  ☐ NO GO','']], [2.4,1.6,2.6])
    return save(d,'07_Checklist_avant_formation_MADGI_ESR.docx')

def rehearsal():
    d=Document(); setup(d,'CONDUCTEUR DE RÉPÉTITION','Préparer l’animation, les transitions et les réponses difficiles','08')
    section(d,'Répétition 1 — Maîtrise du fond (J-3)')
    numbered(d,['Présenter le parcours complet sans lire les supports.','Chronométrer chaque séquence et noter les dépassements.','Exécuter la démonstration avec les données de test.','Rejouer le workflow chèque et verbaliser la différence entre saisi, validé et encaissé.','Vérifier que l’explication du recalcul spontané est comprise par un collègue.'])
    section(d,'Répétition 2 — Conditions réelles (J-1)')
    numbered(d,['Utiliser le même poste, navigateur et vidéoprojecteur que le jour J.','Tester chaque lien, compte et téléchargement.','Simuler une erreur réseau et expliquer la conduite à tenir.','Faire jouer le rôle d’un participant difficile ou pressé.','Terminer en 45 minutes pour conserver le temps des exercices.'])
    section(d,'Formulations clés')
    table(d,['Question utilisateur','Réponse courte et exacte'],[
      ['« Où est passé mon chèque ? »','Il reste dans Validation des paiements avec son statut. Recherchez-le puis poursuivez les étapes bancaires.'],
      ['« Quand mon compte est-il crédité ? »','Uniquement lorsque le paiement atteint le statut ENCAISSÉ.'],
      ['« Le spontané enlève-t-il 500 000 du trimestre ? »','Non. Il augmente le capital acquis, puis le logiciel recalcule la cotisation sur le besoin restant et les trimestres restants.'],
      ['« Pourquoi je ne peux pas clôturer ? »','Le contrôle affiche un blocage ou votre profil n’a pas l’autorisation.'],
      ['« Puis-je recommencer après une erreur ? »','Vérifiez d’abord la liste et le statut pour éviter un doublon, puis contactez le support si nécessaire.']], [2.25,4.35])
    section(d,'Grille d’autoévaluation du formateur')
    table(d,['Critère','1','2','3','4','Commentaire'],[['Respect du temps','','','','',''],['Clarté des explications','','','','',''],['Fluidité de la démonstration','','','','',''],['Gestion des questions','','','','',''],['Maîtrise des cas d’erreur','','','','','']], [2.4,0.45,0.45,0.45,0.45,2.2])
    add_callout(d,'Plan B','Conserver des captures d’écran des étapes essentielles et le scénario imprimé. Si l’environnement est indisponible, présenter le flux, réaliser les exercices sur papier et reprogrammer la pratique applicative.')
    return save(d,'08_Conducteur_repetition_MADGI_ESR.docx')

if __name__=='__main__':
    files=[program(),demo(),guide(),fiches(),evaluation(),checklist(),rehearsal()]
    print('\n'.join(str(p) for p in files))
