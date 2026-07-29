from __future__ import annotations

import json
import re
from pathlib import Path

import pdfplumber
from docx import Document
from openpyxl import load_workbook


ROOT = Path(r"C:\PROJETS\madgi-esr\DOCUMENT ESR")
OUT = Path(r"C:\PROJETS\madgi-esr\.codex-esr-audit")
OUT.mkdir(parents=True, exist_ok=True)


def clean(value: object) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def read_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [
        {"style": p.style.name if p.style else "", "text": clean(p.text)}
        for p in doc.paragraphs
        if clean(p.text)
    ]
    tables = []
    for table in doc.tables:
        tables.append([[clean(cell.text) for cell in row.cells] for row in table.rows])
    return {"type": "docx", "paragraphs": paragraphs, "tables": tables}


def read_pdf(path: Path) -> dict:
    pages = []
    with pdfplumber.open(path) as pdf:
        for number, page in enumerate(pdf.pages, 1):
            pages.append({"page": number, "text": page.extract_text() or ""})
    return {"type": "pdf", "pages": pages}


def read_excel(path: Path) -> dict:
    formula_wb = load_workbook(path, read_only=False, data_only=False, keep_vba=path.suffix.lower() == ".xlsm")
    value_wb = load_workbook(path, read_only=False, data_only=True, keep_vba=path.suffix.lower() == ".xlsm")
    result = {"type": path.suffix.lower().lstrip("."), "sheets": []}
    for ws in formula_wb.worksheets:
        value_ws = value_wb[ws.title]
        nonempty = []
        formulas = []
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    nonempty.append(
                        {
                            "cell": cell.coordinate,
                            "value": clean(cell.value),
                            "computed": clean(value_ws[cell.coordinate].value),
                            "number_format": cell.number_format,
                        }
                    )
                    if isinstance(cell.value, str) and cell.value.startswith("="):
                        formulas.append(
                            {
                                "cell": cell.coordinate,
                                "formula": cell.value,
                                "computed": clean(value_ws[cell.coordinate].value),
                            }
                        )
        result["sheets"].append(
            {
                "name": ws.title,
                "max_row": ws.max_row,
                "max_column": ws.max_column,
                "nonempty": nonempty,
                "formulas": formulas,
            }
        )
    return result


manifest = []
for path in sorted(ROOT.iterdir()):
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            content = read_docx(path)
        elif suffix == ".pdf":
            content = read_pdf(path)
        elif suffix in {".xlsx", ".xlsm"}:
            content = read_excel(path)
        elif suffix == ".doc":
            content = {"type": "legacy_doc", "note": "Extraction Word COM requise"}
        else:
            content = {"type": "unsupported"}
        status = "ok"
    except Exception as exc:
        content = {"type": suffix, "error": repr(exc)}
        status = "error"
    target = OUT / f"{path.name}.json"
    target.write_text(json.dumps(content, ensure_ascii=False, indent=2), encoding="utf-8")
    manifest.append(
        {
            "file": path.name,
            "size": path.stat().st_size,
            "status": status,
            "extraction": str(target),
        }
    )

(OUT / "manifest.json").write_text(
    json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
)
print(json.dumps(manifest, ensure_ascii=False, indent=2))
