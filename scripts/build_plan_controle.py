from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(r"C:\PROJETS\madgi-esr\output\Plan_de_controle_MADGI_ESR.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
PALE = "F4F6F9"
GREEN = "1E6B4E"
AMBER = "9A6700"
RED = "9B1C1C"
GRAY = "5B6573"
WHITE = "FFFFFF"
BLACK = "111827"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(table, color="CBD5E1", size="6"):
    tbl_pr = table._tbl.tblPr
    el = tbl_pr.find(qn("w:tblBorders"))
    if el is None:
        el = OxmlElement("w:tblBorders")
        tbl_pr.append(el)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)
        el.append(tag)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def set_run(run, size=10.5, bold=False, color=BLACK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_p(text="", style=None, bold=False, color=BLACK, size=10.5, align=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    if align is not None:
        p.alignment = align
    return p


def add_heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text), size=10.5)
    return p


def add_number(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text), size=10.5)
    return p


def add_callout(title, text, color=NAVY, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [9360])
    borders(table, "D5DEE8")
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(title + "  "), size=10.5, bold=True, color=color)
    set_run(p.add_run(text), size=10.5, color=BLACK)
    add_p(after=2)


def add_check_table(rows):
    table = doc.add_table(rows=1, cols=5)
    widths = [520, 3160, 2820, 1500, 1360]
    table_geometry(table, widths)
    borders(table)
    headers = ["N°", "Contrôle à effectuer", "Résultat attendu", "Preuve", "Statut"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(text), size=8.5, bold=True, color=WHITE)
    for idx, (control, expected, proof) in enumerate(rows, 1):
        cells = table.add_row().cells
        values = [str(idx), control, expected, proof, "[ ] C  [ ] P\n[ ] NC [ ] BE"]
        for i, text in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(text), size=8.2, color=BLACK)
            if idx % 2 == 0:
                shade(cells[i], "F8FAFC")
        for cell in cells:
            cell_margins(cell, top=110, bottom=110)
    return table


def add_command_block(lines):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [9360])
    borders(table, "9CA3AF")
    cell = table.cell(0, 0)
    shade(cell, "F3F4F6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        run = p.add_run(line + ("" if idx == len(lines) - 1 else "\n"))
        set_run(run, size=8.5, font="Consolas", color="111827")
    add_p(after=2)


def page_break():
    return None


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2
for name, size, color, before, after in (
    ("Heading 1", 16, NAVY, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 11.5, NAVY, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header.add_run("MADGI ESR  |  Plan de contrôle"), size=8.5, bold=True, color=GRAY)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(footer.add_run("MADGI ESR - Document de contrôle"), size=8, color=GRAY)

# Cover
add_p("MADGI", bold=True, color=GREEN, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, before=40, after=8)
add_p("ÉPARGNE SANTÉ RETRAITE", bold=True, color=NAVY, size=17, align=WD_ALIGN_PARAGRAPH.CENTER, after=44)
add_p("PLAN DE CONTRÔLE ET DE VÉRIFICATION", bold=True, color=NAVY, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
add_p("Audit des travaux réalisés selon le plan de finalisation", color=BLUE, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=40)
add_callout(
    "MISSION DU COLLABORATEUR",
    "Vérifier chaque lot sans modifier le code, la configuration ou les données de production. "
    "Toute opération créant des données doit être réalisée uniquement dans un environnement de recette autorisé.",
    color=RED,
    fill="FFF4F2",
)
meta = doc.add_table(rows=6, cols=2)
table_geometry(meta, [2700, 6660])
borders(meta, "D5DEE8")
metadata = [
    ("Contrôleur désigné", "........................................................................"),
    ("Fonction", "........................................................................"),
    ("Date de début", "...................................."),
    ("Date de fin", "...................................."),
    ("Environnement", "[ ] Recette   [ ] Développement   [ ] Production (lecture seule)"),
    ("Version du document", "1.0 - 29 juillet 2026"),
]
for row, (label, value) in zip(meta.rows, metadata):
    shade(row.cells[0], LIGHT_BLUE)
    set_run(row.cells[0].paragraphs[0].add_run(label), size=9, bold=True, color=NAVY)
    set_run(row.cells[1].paragraphs[0].add_run(value), size=9, color=BLACK)
add_p(after=18)
add_p("Document à suivre dans l’ordre. Ne cocher une exigence que si la preuve demandée est disponible.",
      bold=True, color=AMBER, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()
add_heading("1. Objet, périmètre et règles de mission", 1)
add_p(
    "Ce document permet à un collaborateur de contrôler méthodiquement les travaux réalisés sur "
    "MADGI ESR. Il suit les lots du plan de finalisation : gouvernance, paramètres, moteur actuariel, "
    "cotisations, précomptes, paiements, prestations, documents, reporting, sécurité, sauvegarde et production."
)
add_heading("1.1 Interdictions pendant l’audit", 2)
for text in [
    "Ne modifier aucun fichier du projet.",
    "Ne modifier aucun paramètre, rôle, secret ou donnée de production.",
    "Ne créer aucun paiement, précompte ou dossier de prestation en production.",
    "Ne supprimer aucune donnée et ne lancer aucune restauration sur la base active.",
    "Ne communiquer dans le rapport aucun mot de passe, OTP, jeton ou clé technique.",
]:
    add_bullet(text)
add_heading("1.2 Statuts à utiliser", 2)
status_table = doc.add_table(rows=1, cols=2)
table_geometry(status_table, [1900, 7460])
borders(status_table)
for i, h in enumerate(("Statut", "Définition obligatoire")):
    shade(status_table.rows[0].cells[i], NAVY)
    set_run(status_table.rows[0].cells[i].paragraphs[0].add_run(h), size=9, bold=True, color=WHITE)
for code, desc, fill in [
    ("C - Conforme", "Le résultat attendu est obtenu et une preuve vérifiable est jointe.", "EAF7F0"),
    ("P - Partiel", "La fonction existe, mais un contrôle, une preuve ou un sous-cas manque.", "FFF8E5"),
    ("NC - Non conforme", "Le résultat obtenu contredit le résultat attendu.", "FDECEC"),
    ("BE - Bloqué externe", "Le contrôle exige une décision, un accès, un modèle ou une donnée non disponible.", "EEF2F7"),
]:
    cells = status_table.add_row().cells
    shade(cells[0], fill)
    set_run(cells[0].paragraphs[0].add_run(code), size=9, bold=True, color=BLACK)
    set_run(cells[1].paragraphs[0].add_run(desc), size=9, color=BLACK)
add_heading("1.3 Preuves acceptables", 2)
for text in [
    "Capture d’écran datée montrant l’entrée et le résultat.",
    "Fichier PDF ou Excel téléchargé et ouvert avec succès.",
    "Sortie complète d’une commande de test.",
    "Extrait du journal d’audit sans secret ni donnée sensible inutile.",
    "Comparaison chiffrée avec le classeur de référence.",
    "Procès-verbal ou décision signée pour les éléments d’arbitrage.",
]:
    add_bullet(text)

add_heading("2. Préparation de la mission", 1)
add_number("Créer un dossier de preuves nommé CONTROLE_MADGI_ESR_AAAA-MM-JJ.")
add_number("Créer un sous-dossier par lot : 01_Technique, 02_Parametres, 03_Actuariel, etc.")
add_number("Lire le plan de finalisation et la matrice de finalisation avant tout test.")
add_number("Confirmer par écrit l’environnement autorisé pour les tests qui créent des données.")
add_number("Noter l’URL, la date, l’heure et le profil utilisé pour chaque scénario.")
add_callout(
    "ARRÊT IMMÉDIAT",
    "Si un test exige un secret absent, une action destructive, une donnée réelle non anonymisée ou une "
    "modification de production, arrêter le scénario et le classer BE - Bloqué externe.",
    color=RED,
    fill="FFF4F2",
)

page_break()
add_heading("3. Lot 1 - Contrôle technique", 1)
add_p("Objectif : prouver que le frontend et l’API compilent, que les tests passent et que les dépendances ne présentent pas d’alerte connue.")
add_heading("3.1 Procédure", 2)
add_number("Ouvrir PowerShell.")
add_number(r"Se placer dans C:\PROJETS\madgi-esr.")
add_number("Exécuter les commandes ci-dessous sans en modifier le contenu.")
add_command_block([
    r"cd C:\PROJETS\madgi-esr",
    "npm run lint",
    "npm run build",
    "npm audit --omit=dev",
    "cd server",
    "npm run lint",
    "npm test",
    "npm run build",
    "npm audit --omit=dev",
])
add_number("Enregistrer toute la sortie dans le dossier 01_Technique.")
add_check_table([
    ("Validation TypeScript du frontend", "Commande terminée sans erreur.", "Sortie lint"),
    ("Build de production du frontend", "Build réussi ; aucun échec.", "Sortie build"),
    ("Audit des dépendances frontend", "0 vulnérabilité.", "Sortie npm audit"),
    ("Validation TypeScript du backend", "Commande terminée sans erreur.", "Sortie lint"),
    ("Tests automatisés", "19 tests ou davantage, 0 échec.", "Sortie npm test"),
    ("Build de production du backend", "Build réussi ; aucun échec.", "Sortie build"),
    ("Audit des dépendances backend", "0 vulnérabilité.", "Sortie npm audit"),
])
add_callout("NOTE", "L’avertissement sur la taille des fichiers JavaScript n’est pas un échec de build. Le consigner comme observation de performance.", color=AMBER, fill="FFF8E5")

page_break()
add_heading("4. Lot 2 - Paramètres actuariels", 1)
add_p("Objectif : vérifier que les valeurs métier sont administrables et datées, sans modifier leur valeur pendant l’audit.")
add_heading("4.1 Procédure", 2)
add_number("Se connecter avec un profil Gestionnaire ou Administrateur autorisé.")
add_number("Ouvrir le menu Paramètres, puis Paramètres généraux ESR.")
add_number("Pour chaque ligne, relever le code, la valeur, la date de début, la date de fin et l’état actif.")
add_number("Ne cliquer sur Enregistrer sous aucun prétexte pendant ce contrôle en lecture seule.")
add_check_table([
    ("TAUX_GAR", "Valeur numérique active et datée.", "Capture"),
    ("FRAIS_RENTE", "Valeur numérique active et datée.", "Capture"),
    ("FRAIS_GESTION_RACHAT", "Valeur numérique active et datée.", "Capture"),
    ("TAUX_RACHAT", "Valeur numérique active et datée.", "Capture"),
    ("DELAI_MIN_RACHAT_ANNEES", "Valeur numérique active et datée.", "Capture"),
    ("TAUX_DECES_AVANT_RETRAITE", "Valeur numérique active et datée.", "Capture"),
    ("TAUX_INVALIDITE_AVANT_RETRAITE", "Valeur numérique active et datée.", "Capture"),
    ("TAUX_COUVERTURE_RETRAITE", "Valeur numérique active et datée.", "Capture"),
    ("TAUX_REMBOURSEMENT_SOINS", "Valeur numérique active et datée.", "Capture"),
    ("TAUX_DECES_PENDANT_RENTE", "Valeur numérique active et datée.", "Capture"),
    ("AGE_MAX et table CIMA F", "Âge maximum cohérent ; âges et lx présents.", "Capture/export"),
])
add_callout("CLASSEMENT", "Une valeur présente mais non validée officiellement est P - Partiel ou BE - Bloqué externe, jamais C - Conforme métier.", color=AMBER, fill="FFF8E5")

page_break()
add_heading("5. Lot 3 - Moteur actuariel et comptes ESR", 1)
add_p("Objectif : rapprocher les résultats du moteur avec les classeurs de référence. Tolérance maximale : 0,1 %.")
add_heading("5.1 Règle de comparaison", 2)
add_p("Calculer l’écart relatif avec la formule suivante :")
add_command_block(["Écart relatif (%) = |Résultat application - Résultat de référence| / Résultat de référence × 100"])
add_p("Si le résultat de référence est nul, comparer les montants directement et expliquer l’écart.")
add_heading("5.2 Scénarios actuariels", 2)
add_check_table([
    ("Cotisation trimestrielle - cas 60 ans", "Écart ≤ 0,1 % avec le classeur.", "Tableau comparatif"),
    ("Cotisation trimestrielle - cas 65 ans", "Écart ≤ 0,1 % avec le classeur.", "Tableau comparatif"),
    ("Cotisation unique à la retraite", "PU égale au capital constitutif.", "Capture/calcul"),
    ("Cotisation unique avant retraite", "PU actualisée inférieure au capital constitutif.", "Capture/calcul"),
    ("Provision mathématique individuelle", "Résultat conforme au classeur PM.", "Tableau comparatif"),
    ("Rachat après délai minimum", "Frais et pénalité appliqués selon paramètres.", "Liquidation/calcul"),
    ("Rachat avant délai minimum", "Opération refusée et tracée.", "Message/audit"),
    ("Décès avant retraite", "Taux paramétré appliqué à la provision.", "Calcul/audit"),
    ("Invalidité", "Taux paramétré appliqué à la provision.", "Calcul/audit"),
    ("Décès pendant rente", "Taux paramétré appliqué au capital restant.", "Calcul/audit"),
])
add_heading("5.3 Compte individuel", 2)
add_number("Ouvrir Comptes ESR et sélectionner un adhérent disposant de mouvements.")
add_number("Relever le capital acquis, la PM, la valeur de rachat, la date et la version.")
add_number("Comparer le capital avec la somme des seules cotisations encaissées.")
add_number("Télécharger l’avis annuel et contrôler les mêmes montants.")
add_check_table([
    ("Cotisations encaissées seulement", "Les mouvements prévus/rejetés n’alimentent pas le compte.", "Liste + compte"),
    ("Capital acquis", "Somme des montants encaissés à la date.", "Calcul"),
    ("Provision mathématique", "Capitalisation au taux daté.", "Calcul"),
    ("Valeur de rachat", "Frais/pénalité paramétrés appliqués.", "Calcul"),
    ("Traçabilité", "Date et version du calcul affichées.", "Capture"),
    ("Avis annuel", "PDF lisible et montants identiques.", "PDF"),
])

page_break()
add_heading("6. Lot 4 - Précomptes et retour DGI", 1)
add_callout("ENVIRONNEMENT", "Les imports et générations doivent être réalisés uniquement en recette. En production, limiter le contrôle à la consultation.", color=RED, fill="FFF4F2")
add_heading("6.1 Génération", 2)
add_number("Choisir une période de recette autorisée.")
add_number("Générer le précompte une première fois et relever les volumes et montants.")
add_number("Relancer la même génération.")
add_number("Vérifier qu’aucun doublon n’est créé.")
add_heading("6.2 Retour DGI", 2)
add_number("Préparer trois lignes : montant identique, montant différent et montant nul.")
add_number("Importer le fichier retour et relever le résumé affiché.")
add_number("Vérifier les cotisations et l’état des non-précomptés.")
add_check_table([
    ("Éligibilité des adhérents", "Seuls les adhérents éligibles sont générés.", "Export/capture"),
    ("Prévention des doublons", "Deuxième exécution sans nouveau doublon.", "Comparaison"),
    ("Retour identique", "Statut ENCAISSE.", "Capture"),
    ("Retour différent", "Statut ECART ; montant reçu conservé.", "Capture"),
    ("Retour nul", "Statut NON_PRECOMPTE.", "Capture"),
    ("Motif", "Motif présent dans la traçabilité d’import.", "Audit"),
    ("État non-précomptés", "Matricule, téléphone, montant et statut disponibles.", "Capture/export"),
    ("Fichier réel APS/DGI", "Aller-retour homologué sur format officiel.", "PV APS/DGI"),
])
add_callout("ATTENTION", "Sans fichier officiel APS/DGI, le dernier contrôle doit rester BE - Bloqué externe.", color=AMBER, fill="FFF8E5")

page_break()
add_heading("7. Lot 5 - Paiements spontanés", 1)
add_p("Objectif : vérifier la séparation des responsabilités et l’alimentation unique du compte à l’encaissement.")
add_heading("7.1 Scénario nominal en recette", 2)
for text in [
    "Créer un paiement de test : le statut initial doit être SAISI.",
    "Passer à CONTROLE.",
    "Passer à VALIDE.",
    "Passer à ENCAISSE.",
    "Télécharger le reçu.",
    "Vérifier qu’une seule cotisation encaissée a été créée.",
]:
    add_number(text)
add_heading("7.2 Scénarios de refus", 2)
add_check_table([
    ("SAISI vers ENCAISSE", "Transition refusée.", "Message"),
    ("CONTROLE vers ENCAISSE", "Transition refusée.", "Message"),
    ("Paiement REJETE", "Aucune transition ultérieure.", "Message"),
    ("Paiement ENCAISSE", "Aucune modification ultérieure.", "Message"),
    ("Reçu avant encaissement", "Téléchargement refusé.", "Message"),
    ("Reçu après encaissement", "PDF disponible et exact.", "PDF"),
    ("Compte ESR", "Une seule cotisation créée à l’encaissement.", "Compte/liste"),
    ("Audit", "Chaque transition contient auteur, date et observation.", "Journal"),
])

page_break()
add_heading("8. Lot 6 - Prestations, rachats et invalidité", 1)
add_p("Objectif : vérifier que le montant n’est pas saisi librement et que le workflow est respecté.")
add_heading("8.1 Circuit obligatoire", 2)
add_command_block(["DOSSIER_OUVERT  →  EN_CONTROLE  →  VALIDE  →  PAYE"])
add_number("En recette, ouvrir un dossier pour un adhérent disposant d’un compte calculé.")
add_number("Vérifier que le formulaire ne permet pas de saisir librement le montant.")
add_number("Passer successivement par les trois transitions autorisées.")
add_number("Télécharger la liquidation et contrôler le montant.")
add_number("Consulter le journal d’audit.")
add_check_table([
    ("Montant automatique", "Montant issu du compte et des paramètres datés.", "Écran/audit"),
    ("Saut OUVERT vers VALIDE", "Transition refusée.", "Message"),
    ("Saut OUVERT vers PAYE", "Transition refusée.", "Message"),
    ("Dossier en contrôle", "Date de complétude et échéance calculées.", "Audit"),
    ("Délai de 15 jours", "Week-ends exclus du calcul.", "Calcul"),
    ("Dossier PAYE", "Statut terminal.", "Message"),
    ("Rachat partiel", "Option absente ou explicitement refusée.", "Écran/message"),
    ("Rachat avant délai", "Refus et événement d’audit.", "Message/audit"),
    ("Liquidation PDF", "Identité, type, montant et statut exacts.", "PDF"),
    ("Pièces justificatives", "Liste et complétude validées selon procédure.", "Dossier/PV"),
])
add_callout("RÈGLES NON SIGNÉES", "Les taux 95/100 %, l’arrêt de rente et les pièces médicales doivent rester P ou BE jusqu’à décision signée.", color=AMBER, fill="FFF8E5")

page_break()
add_heading("9. Lot 7 - Rentes", 1)
add_p("Objectif : distinguer les fonctions consultables des règles qui attendent encore un arbitrage.")
add_check_table([
    ("Liste des rentes", "Rentes chargées sans erreur.", "Capture"),
    ("Identité du rentier", "Matricule, nom et prénoms exacts.", "Capture"),
    ("Capital initial et restant", "Montants cohérents avec la liquidation.", "Comparaison"),
    ("Versements", "Échéances rattachées à la bonne rente.", "Capture/export"),
    ("Décès pendant rente", "80 % ou taux paramétré calculé sur capital restant.", "Calcul"),
    ("Création et échéancier", "Processus de bout en bout validé.", "PV"),
    ("Suspension et extinction", "Règles signées et scénarios réussis.", "PV"),
    ("Bordereau APS", "Format officiel accepté.", "Fichier/PV APS"),
])
add_callout("CLASSEMENT ATTENDU", "Création, suspension, extinction et bordereau APS sont BE tant que les règles et formats externes ne sont pas fournis.", color=AMBER, fill="FFF8E5")

page_break()
add_heading("10. Lot 8 - Documents et reporting", 1)
add_heading("10.1 Contrôle des PDF", 2)
add_number("Télécharger chaque document depuis l’application.")
add_number("L’ouvrir à 100 % de zoom.")
add_number("Vérifier l’absence de texte coupé, superposé ou illisible.")
add_number("Comparer identité, dates et montants avec l’écran source.")
add_check_table([
    ("Reçu de paiement", "Disponible après encaissement uniquement.", "PDF"),
    ("Liquidation", "Montant calculé et identité exacts.", "PDF"),
    ("Avis annuel", "Capital, PM et rachat identiques au compte.", "PDF"),
    ("Numérotation", "Référence unique et lisible.", "PDF"),
    ("Lisibilité", "Aucune coupure, superposition ou caractère corrompu.", "Capture"),
    ("Archivage", "Document rattachable à l’opération.", "Dossier/PV"),
])
add_heading("10.2 Reporting CIMA C-20", 2)
add_number("Depuis le Tableau de bord, lancer l’export CIMA C-20 Excel.")
add_number("Ouvrir le classeur et vérifier qu’Excel ne signale aucune réparation.")
add_number("Comparer les cotisations encaissées, le capital, la PM et les prestations avec la base.")
add_check_table([
    ("Ouverture Excel", "Aucun avertissement ni réparation.", "Classeur"),
    ("Cotisations par trimestre", "Totaux réconciliés avec les mouvements.", "Tableau"),
    ("Provision globale", "Somme des PM individuelles.", "Tableau"),
    ("Capital acquis", "Somme des comptes individuels.", "Tableau"),
    ("Prestations payées", "Somme des montants payés.", "Tableau"),
    ("Modèle réglementaire", "Présentation homologuée par Finance/ESR.", "PV signé"),
])

page_break()
add_heading("11. Lot 9 - Profils, droits et audit", 1)
add_p("Exécuter les contrôles avec un compte distinct pour chaque profil. Ne jamais modifier le rôle d’un compte de production pendant l’audit.")
add_check_table([
    ("Adhérent - données personnelles", "Accès à son dossier et son compte seulement.", "Captures"),
    ("Adhérent - administration", "Accès refusé aux paramètres, audit et utilisateurs.", "Captures"),
    ("Gestionnaire", "Accès aux opérations ; pas d’administration globale.", "Captures"),
    ("Administrateur", "Accès aux utilisateurs, paramètres et audit.", "Captures"),
    ("Superadministrateur", "Accès exceptionnel conforme à la matrice.", "Captures"),
    ("Compte désactivé", "Connexion ou accès refusé.", "Capture"),
    ("Audit des calculs", "Entrées, résultat, paramètres et auteur présents.", "Journal"),
    ("Audit des paiements", "Chaque transition présente.", "Journal"),
    ("Audit des prestations", "Création, refus et transitions présents.", "Journal"),
    ("Audit des paramètres", "Modification sensible traçable.", "Journal"),
])
add_heading("11.1 Contrôle des secrets", 2)
for text in [
    "Aucune clé service_role ne doit apparaître dans le navigateur.",
    "Aucun fichier .env réel ne doit être joint au rapport.",
    "Un OTP fixe doit empêcher le démarrage en mode production.",
    "HTTPS, domaine et certificats doivent être contrôlés par la DSI sur l’environnement cible.",
]:
    add_bullet(text)

page_break()
add_heading("12. Lot 10 - Sauvegarde, reprise et production", 1)
add_heading("12.1 Vérification non destructive de sauvegarde", 2)
add_command_block([
    r"cd C:\PROJETS\madgi-esr",
    "npm run backup:data",
    'node scripts/verify-backup.mjs "CHEMIN_DU_DOSSIER_AFFICHE"',
])
add_check_table([
    ("Création de sauvegarde", "Un dossier horodaté est créé.", "Dossier"),
    ("Manifeste", "Date, source, tables, volumes et SHA-256 présents.", "manifest.json"),
    ("Intégrité", "ok: true.", "Sortie commande"),
    ("Tables", "15 tables vérifiées ou davantage.", "Sortie commande"),
    ("Relations", "10 relations vérifiées ou davantage.", "Sortie commande"),
    ("Copie hors plateforme", "Stockage chiffré distinct confirmé.", "PV DSI"),
    ("Restauration isolée", "Base restaurée et contrôles fonctionnels réussis.", "Rapport DSI"),
    ("Retour arrière", "Procédure validée et chronométrée.", "PV DSI"),
])
add_callout("INTERDICTION", "Ne jamais tester une restauration sur la base active. Sans environnement isolé fourni par la DSI, classer BE.", color=RED, fill="FFF4F2")

add_heading("13. Contrôles externes et décisions", 1)
add_check_table([
    ("Rachat D01-D02", "Règle, frais, pénalité, arrondi et date de valeur signés.", "PV signé"),
    ("Décès D03", "Taux et base signés.", "PV signé"),
    ("Invalidité D04", "Taux, pièces et délai signés.", "PV signé"),
    ("Couverture D05", "Soins et cotisation maladie distingués.", "PV signé"),
    ("Décès pendant rente D06", "Date d’arrêt et trimestre du décès signés.", "PV signé"),
    ("Architecture D07", "React/Express/Supabase homologué ou réécriture décidée.", "Décision DSI"),
    ("APS/DGI/SIAPS", "Accès, dictionnaire et formats homologués.", "PV technique"),
    ("Historique", "Données détaillées fournies et totaux signés.", "PV Finance"),
    ("Placements", "Périmètre V1/V2 décidé.", "Décision Direction"),
    ("Production", "Domaine, secrets, sauvegardes et fenêtre validés.", "PV DSI"),
])

page_break()
add_heading("14. Fiche d’anomalie", 1)
add_p("Créer une fiche par anomalie. Ne jamais regrouper plusieurs causes indépendantes dans la même fiche.")
anomaly = doc.add_table(rows=10, cols=2)
table_geometry(anomaly, [2500, 6860])
borders(anomaly)
fields = [
    ("Référence", "ANO-____"),
    ("Lot / contrôle", ""),
    ("Date et profil", ""),
    ("Environnement", ""),
    ("Étapes de reproduction", "\n\n"),
    ("Résultat attendu", "\n"),
    ("Résultat obtenu", "\n"),
    ("Preuve jointe", ""),
    ("Criticité", "[ ] Bloquante  [ ] Majeure  [ ] Mineure"),
    ("Statut", "[ ] Ouverte  [ ] Corrigée  [ ] Re-testée  [ ] Acceptée"),
]
for row, (label, value) in zip(anomaly.rows, fields):
    shade(row.cells[0], LIGHT_BLUE)
    set_run(row.cells[0].paragraphs[0].add_run(label), size=9, bold=True, color=NAVY)
    set_run(row.cells[1].paragraphs[0].add_run(value), size=9, color=BLACK)

add_heading("15. Synthèse finale", 1)
summary = doc.add_table(rows=1, cols=6)
table_geometry(summary, [700, 3420, 1100, 1100, 1100, 1940])
borders(summary)
for i, h in enumerate(("N°", "Lot", "C", "P", "NC", "BE / Observation")):
    shade(summary.rows[0].cells[i], NAVY)
    set_run(summary.rows[0].cells[i].paragraphs[0].add_run(h), size=8.5, bold=True, color=WHITE)
lots = [
    "Contrôle technique", "Paramètres", "Moteur actuariel et comptes", "Précomptes/DGI",
    "Paiements", "Prestations", "Rentes", "Documents/reporting", "Profils/sécurité",
    "Sauvegarde/reprise", "Décisions externes",
]
for idx, lot in enumerate(lots, 1):
    cells = summary.add_row().cells
    for i, value in enumerate((str(idx), lot, "", "", "", "")):
        set_run(cells[i].paragraphs[0].add_run(value), size=8.5, color=BLACK)
        if idx % 2 == 0:
            shade(cells[i], "F8FAFC")

p16 = add_heading("16. Conclusion du contrôleur", 1)
p16.paragraph_format.page_break_before = True
conclusion = doc.add_table(rows=6, cols=2)
table_geometry(conclusion, [2800, 6560])
borders(conclusion)
conclusion_fields = [
    ("Décision proposée", "[ ] Conforme pour recette\n[ ] Conforme sous réserves\n[ ] Non conforme\n[ ] Bloqué externe"),
    ("Anomalies bloquantes", "\n\n"),
    ("Réserves", "\n\n"),
    ("Actions demandées", "\n\n"),
    ("Contrôleur / signature", "\n"),
    ("Responsable ESR / visa", "\n"),
]
for row, (label, value) in zip(conclusion.rows, conclusion_fields):
    shade(row.cells[0], LIGHT_BLUE)
    set_run(row.cells[0].paragraphs[0].add_run(label), size=9, bold=True, color=NAVY)
    set_run(row.cells[1].paragraphs[0].add_run(value), size=9, color=BLACK)

add_p(after=12)
add_callout(
    "RÈGLE DE CLÔTURE",
    "Le projet ne peut être déclaré prêt pour la production que si tous les contrôles bloquants sont C - Conforme, "
    "que les anomalies bloquantes sont corrigées et re-testées, et que les validations externes sont signées.",
    color=RED,
    fill="FFF4F2",
)

doc.save(OUT)
print(OUT)
