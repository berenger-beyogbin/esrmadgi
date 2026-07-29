from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\PROJETS\madgi-esr")
OUTPUT = ROOT / "output" / "Plan_de_finalisation_MADGI_ESR.docx"
LOGO = ROOT / "src" / "assets" / "logos" / "logo-madgi.jpg"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FA"
GREEN = "2E7D32"
LIGHT_GREEN = "E2F0D9"
AMBER = "B26A00"
LIGHT_AMBER = "FFF2CC"
RED = "C00000"
LIGHT_RED = "FCE4D6"
GRAY = "666666"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"
BLACK = "000000"

STATUS = {
    "Disponible": (GREEN, LIGHT_GREEN),
    "Partiel": (AMBER, LIGHT_AMBER),
    "Absent": (RED, LIGHT_RED),
    "Arbitrage": (NAVY, LIGHT_BLUE),
}

PRIORITY = {
    "Bloquante": (RED, LIGHT_RED),
    "Haute": (AMBER, LIGHT_AMBER),
    "Moyenne": (DARK_BLUE, LIGHT_BLUE),
    "Fin de projet": (GRAY, LIGHT_GRAY),
    "Décision": (NAVY, LIGHT_BLUE),
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_inches: list[float], indent_dxa=120) -> None:
    widths_dxa = [round(width * 1440) for width in widths_inches]
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_inches[index])
            set_cell_margins(cell)


def set_run_font(run, size=11, color=BLACK, bold=False, italic=False, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_text(cell, text: str, *, bold=False, color=BLACK, size=8.4,
                  align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=8.5, color=GRAY)


def add_paragraph(doc, text: str, *, size=11, color=BLACK, bold=False,
                  italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=6,
                  before=0, keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_with_next = keep_with_next
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text: str, *, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    # Empty style-created run is harmless; explicit run guarantees the token map.
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
    return p


def add_callout(doc, label: str, text: str, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"{label} ")
    set_run_font(r1, size=10.5, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float],
              *, font_size=8.4, status_col: int | None = None,
              priority_col: int | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        set_cell_shading(header.cells[i], LIGHT_BLUE)
        set_cell_text(header.cells[i], text, bold=True, color=NAVY, size=8.5,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in {status_col, priority_col} else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[i], value, size=font_size, align=align)
            if status_col is not None and i == status_col and value in STATUS:
                color, fill = STATUS[value]
                set_cell_shading(row.cells[i], fill)
                set_cell_text(row.cells[i], value.upper(), bold=True, color=color,
                              size=7.8, align=WD_ALIGN_PARAGRAPH.CENTER)
            if priority_col is not None and i == priority_col and value in PRIORITY:
                color, fill = PRIORITY[value]
                set_cell_shading(row.cells[i], fill)
                set_cell_text(row.cells[i], value.upper(), bold=True, color=color,
                              size=7.6, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


comparison_sections = [
    (
        "A. Gouvernance, architecture et accès",
        [
            ["Arbitrage métier", "Notes, contrats et classeurs reçus.", "Signer les règles : rachat, décès, invalidité, couverture et arrondis.", "Arbitrage", "Bloquante"],
            ["Architecture", "React, Express/TypeScript, Supabase/PostgreSQL, connecteurs MySQL/SIAPS.", "Faire approuver l’écart avec l’exigence WebDev/HFSQL.", "Arbitrage", "Bloquante"],
            ["Authentification", "Sessions, première connexion, réinitialisation et démo limitée au développement.", "Recette avec comptes réels et contrôle de configuration production.", "Partiel", "Haute"],
            ["Profils et droits", "Adhérent, Gestionnaire, Administrateur et Superadmin.", "Valider et tester la matrice des autorisations par action.", "Partiel", "Haute"],
            ["Utilisateurs", "Création, modification et gestion des accès.", "Tester activation, désactivation, réactivation et traçabilité.", "Partiel", "Moyenne"],
            ["APS/SIAPS", "Recherche par matricule avec fallback MySQL/SIAPS.", "Recevoir accès, dictionnaire et données de test ; homologuer le connecteur.", "Partiel", "Bloquante"],
            ["Audit", "Page, API et journal d’audit disponibles.", "Couvrir calculs, validations, paiements et paramètres sensibles.", "Partiel", "Haute"],
        ],
    ),
    (
        "B. Adhérents, bénéficiaires et paramètres",
        [
            ["Adhérents", "Création, modification, recherche, statuts et adhésion en ligne.", "Contrôles doublons, historique complet et validations sensibles.", "Partiel", "Haute"],
            ["Bénéficiaires", "CRUD et contrôle de la répartition.", "Éligibilité, historisation et gel au départ à la retraite.", "Partiel", "Haute"],
            ["Bulletin d’adhésion", "Formulaire numérique et modèle papier reçus.", "PDF officiel, numéro, signatures et archivage.", "Partiel", "Moyenne"],
            ["Paramètres actuariels", "Taux, frais, âge maximal, grades et mortalité paramétrables.", "Versionner avec date d’effet et interdire les rétroactivités non autorisées.", "Partiel", "Haute"],
            ["Table CIMA F", "Table reçue et écran d’administration existant.", "Contrôler les coefficients, verrouiller la version et tester les bornes.", "Partiel", "Haute"],
            ["Âge de retraite", "Calcul par grade et date de retraite.", "Faire valider la grille officielle et gérer ses évolutions.", "Partiel", "Haute"],
            ["Simulation", "Calcul affiché dans le formulaire adhérent.", "Créer l’état imprimable officiel et archiver la version calculée.", "Partiel", "Moyenne"],
        ],
    ),
    (
        "C. Moteur actuariel et comptes ESR",
        [
            ["Cotisation trimestrielle", "Moteur actuariel détaillé et script de comparaison.", "Valider tous les cas reçus et retirer les anciens calculs provisoires.", "Partiel", "Haute"],
            ["Cotisation unique", "Formule présente dans la note technique.", "Implémenter calcul, interface, stockage et tests.", "Absent", "Haute"],
            ["Capital constitutif", "Calcul inclus dans le moteur de cotisation.", "Versionner les paramètres et enregistrer le détail de calcul.", "Partiel", "Haute"],
            ["Provision individuelle", "Placeholder dans le code ; exemples Excel reçus.", "Implémenter la formule officielle par adhérent et date d’évaluation.", "Partiel", "Bloquante"],
            ["Provision globale", "Exemples agrégés 2021–2024 reçus.", "Calcul périodique, réconciliation et état comptable.", "Absent", "Haute"],
            ["Compte individuel ESR", "Consultation du capital, PM et valeur de rachat.", "Automatiser l’alimentation et le recalcul à chaque mouvement.", "Partiel", "Haute"],
            ["Révision de cotisation", "Règles décrites dans la note technique.", "Recalculer après modification de R et conserver les versions.", "Absent", "Haute"],
            ["Changement d’âge", "Date de retraite recalculable.", "Recalculer le contrat après changement de grade et historiser.", "Partiel", "Haute"],
            ["Historique de calcul", "Quelques dates et paramètres stockés.", "Conserver formule, version, entrées, résultat, auteur et horodatage.", "Absent", "Haute"],
            ["Tests actuariels", "Un script existe pour la cotisation.", "Matrice signée couvrant PM, rachat, décès, invalidité et rente.", "Partiel", "Bloquante"],
        ],
    ),
    (
        "D. Cotisations, précomptes et paiements",
        [
            ["Cotisations périodiques", "Liste et génération des échéances.", "Capitalisation, régularisations, annulations et périodes incomplètes.", "Partiel", "Haute"],
            ["Paiement spontané", "Formulaire et API d’enregistrement.", "Date de valeur, validation, compte ESR, annulation et reçu.", "Partiel", "Haute"],
            ["Reçu", "Aucun reçu PDF officiel.", "Numéro unique, modèle validé, contrôle et archivage.", "Absent", "Moyenne"],
            ["Paiement par chèque", "Moyen de paiement disponible à la saisie.", "Workflow saisi–contrôlé–validé–rejeté–encaissé.", "Partiel", "Haute"],
            ["Génération précomptes", "Génération trimestrielle et prévention partielle des doublons.", "Valider l’éligibilité et les règles de régularisation.", "Partiel", "Haute"],
            ["Export APS", "Emplacement prévu dans l’écran.", "Bouton actuellement désactivé ; adapter au format officiel APS.", "Absent", "Bloquante"],
            ["Import APS/Comptabilité", "Lecture Excel et enrichissement par matricule.", "Adapter aux colonnes réelles et enregistrer le résultat en base.", "Partiel", "Bloquante"],
            ["Retour DGI", "Rapprochement visuel partiel.", "Persister réalisé, écarts, motifs et régularisations.", "Absent", "Haute"],
            ["Non-précomptés", "Données de base disponibles.", "État dédié avec téléphone, montant et motif.", "Absent", "Moyenne"],
        ],
    ),
    (
        "E. Prestations, rachats, invalidité et rentes",
        [
            ["Dossiers prestations", "Liste et création avec statuts généraux.", "Pièces, transitions, doubles validations, paiement et clôture.", "Partiel", "Haute"],
            ["Rachat total", "Type RACHAT, valeur affichée et classeur reçu.", "Éligibilité, calcul, liquidation, validation et paiement.", "Partiel", "Bloquante"],
            ["Rachat partiel", "Interdit par les conditions générales.", "Bloquer explicitement l’opération et tracer le refus.", "Absent", "Moyenne"],
            ["Décès avant retraite", "Déclaration, prestation et classeur de calcul.", "Calcul 95 %/100 %, pièces, validation et paiement.", "Partiel", "Haute"],
            ["Invalidité totale", "Type de prestation disponible.", "Validation médicale, calcul, pièces et délai de paiement.", "Partiel", "Haute"],
            ["Décès pendant rente", "Règle de 80 % documentée.", "Capital restant, arrêt de rente et répartition bénéficiaires.", "Absent", "Haute"],
            ["Rentes", "Consultation des rentes et versements.", "Création, échéancier, paiement APS, suspension et extinction.", "Partiel", "Haute"],
            ["Versements APS", "Processus métier décrit.", "Bordereau trimestriel, validation et rapprochement.", "Absent", "Haute"],
            ["Délai 15 jours", "Date de demande enregistrée.", "Date de complétude, échéance, alertes et KPI.", "Absent", "Moyenne"],
            ["Placements", "Gouvernance et principes documentés.", "Décider : suivi V1 ou gestion complète V2.", "Arbitrage", "Décision"],
        ],
    ),
    (
        "F. Reporting, données, sécurité et déploiement",
        [
            ["Dashboard", "Indicateurs globaux et dernières opérations.", "Ajouter KPI actuariels, précomptes, délais, écarts et alertes.", "Partiel", "Moyenne"],
            ["Reporting réglementaire", "Exemples CIMA C-20 et PM reçus.", "Produire les états depuis la base sans saisies manuelles.", "Partiel", "Haute"],
            ["Avis annuel", "Espace adhérent imprimable.", "PDF officiel : cotisations, capital, PM, rachat et paramètres.", "Partiel", "Haute"],
            ["Exports PDF", "Impression navigateur disponible.", "Génération serveur stable, numérotation et archivage.", "Partiel", "Moyenne"],
            ["Exports Excel", "Exports ponctuels de précompte.", "États techniques, comptables et modèles validés.", "Partiel", "Moyenne"],
            ["Migration historique", "Classeur 2021–2024 reçu.", "Nettoyer, mapper, importer et réconcilier.", "Absent", "Haute"],
            ["Schéma versionné", "Base Supabase utilisée.", "Ajouter migrations SQL et documentation du modèle au dépôt.", "Absent", "Haute"],
            ["Tests logiciels", "Compilation TypeScript front et back réussie.", "Tests unitaires, API, intégration, sécurité et non-régression.", "Partiel", "Haute"],
            ["Sécurité", "Rôles, Helmet, CORS, rate limit et sessions.", "Audit, secrets, HTTPS, droits et conservation des données.", "Partiel", "Haute"],
            ["Sauvegarde/PRA", "Exigences documentées.", "Sauvegardes, restauration testée, supervision et reprise.", "Absent", "Haute"],
            ["Documentation", "README backend partiel.", "README projet, architecture, exploitation et procédures.", "Partiel", "Moyenne"],
            ["Guides et formation", "Prévus dans le rapport.", "Guides par profil, supports et procès-verbal de formation.", "Absent", "Fin de projet"],
            ["Production", "Configuration Render et exemples d’environnement.", "Domaine, certificats, secrets, sauvegardes et retour arrière.", "Partiel", "Haute"],
        ],
    ),
]


roadmap = [
    ["1", "Arbitrage et gel", "Semaine 1", "Chef ESR + Actuaire", "PV d’arbitrage signé ; architecture approuvée."],
    ["2", "Moteur actuariel", "Semaines 1–2", "Développeur + Actuaire", "Cas tests conformes ; écart ≤ 0,1 %."],
    ["3", "Processus opérationnels", "Semaine 3", "Développeur + ESR + APS", "Précomptes, paiements et prestations de bout en bout."],
    ["4", "Documents et reporting", "Semaine 4", "Développeur + ESR + Finance", "Modèles PDF/Excel validés et réconciliés."],
    ["5", "Données, intégration, sécurité", "Semaine 5", "Développeur + DSI", "Migration contrôlée ; sauvegarde/restauration réussie."],
    ["6", "Recette et production", "Semaine 6", "MADGI + DSI + utilisateurs", "PV de recette, formation et mise en service."],
]

arbitrations = [
    ["Rachat", "Après 2 ans ? Retenue totale 5 % ou 10 % ? Résiliation identique au rachat ?", "Chef ESR + Actuaire", "S1", "À décider"],
    ["Décès avant retraite", "Versement de 95 % ou 100 % de l’épargne constituée ?", "Chef ESR + Actuaire", "S1", "À décider"],
    ["Décès pendant rente", "Date d’arrêt, trimestre du décès et calcul du capital restant dû.", "Actuaire", "S1", "À décider"],
    ["Couverture retraite", "Séparer remboursement des soins à 80 % et paiement de la cotisation maladie.", "Chef ESR", "S1", "À décider"],
    ["Âge de retraite", "Grille officielle par grade et règles d’évolution.", "RH/ESR", "S1", "À décider"],
    ["Architecture", "Maintien React/Express/Supabase ou réécriture WebDev/HFSQL.", "Direction + DSI", "S1", "À décider"],
    ["Placement", "Suivi simple en V1 ou gestion complète dans une V2.", "Direction + Finance", "S1", "À décider"],
]

missing_docs = [
    "Fichier réel de précompte transmis à l’APS.",
    "Fichier retour DGI avec les montants effectivement précomptés.",
    "Dictionnaire de données et accès de test APS/SIAPS.",
    "Modèle officiel d’impression d’une simulation.",
    "Modèle officiel du reçu de paiement spontané.",
    "Workflow validé des paiements par chèque.",
    "Liste définitive et modèles des reportings attendus.",
    "Fiches techniques et état de suivi des placements.",
    "Jeux de tests actuariels signés avec résultats attendus.",
    "Données historiques réelles à migrer.",
    "Décision écrite sur l’hébergement et la technologie.",
]

risks = [
    ["Règle actuarielle erronée", "Très élevé", "Très élevée", "PV signé + tests croisés obligatoires.", "Chef ESR / Actuaire"],
    ["Architecture non approuvée", "Très élevé", "Élevée", "Décision DSI avant poursuite de la finalisation.", "Direction / DSI"],
    ["Format APS indisponible", "Élevé", "Élevée", "Obtenir fichiers aller/retour et données anonymisées.", "APS"],
    ["Historique de mauvaise qualité", "Élevé", "Moyenne", "Profilage, nettoyage, doublons et réconciliation.", "ESR / Développeur"],
    ["Paiement non contrôlé", "Élevé", "Moyenne", "Séparation saisie/validation/encaissement.", "Finance / ESR"],
    ["Absence de restauration testée", "Très élevé", "Moyenne", "Test PRA avant mise en service.", "DSI"],
]

raci = [
    ["Arbitrages métier", "Chef ESR", "Direction", "Actuaire, Finance", "Développeur"],
    ["Formules et cas tests", "Actuaire", "Chef ESR", "Développeur", "Direction"],
    ["Intégration APS/SIAPS", "DSI / APS", "Direction", "Développeur", "Chef ESR"],
    ["Développement", "Développeur", "Chef ESR", "Actuaire, utilisateurs", "Direction"],
    ["Reporting comptable", "Finance", "Chef ESR", "Actuaire, développeur", "Direction"],
    ["Recette métier", "Utilisateurs ESR", "Chef ESR", "Actuaire, DSI", "Direction"],
    ["Mise en production", "DSI", "Direction", "Développeur, ESR", "Utilisateurs"],
]


def build_document() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("MADGI — ÉPARGNE SANTÉ RETRAITE  |  BOUSSOLE DE FINALISATION")
    set_run_font(hr, size=8.5, color=GRAY, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run("Document de pilotage interne  •  25 juillet 2026  •  Page ")
    set_run_font(fr, size=8.5, color=GRAY)
    add_page_field(fp)

    # Cover: editorial-style opening adapted for an operator guide.
    add_paragraph(doc, "DOCUMENT DE PILOTAGE", size=10, color=BLUE, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(18)
        p_logo.add_run().add_picture(str(LOGO), width=Inches(1.05))
    add_paragraph(doc, "PLAN DE FINALISATION", size=27, color=NAVY, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_paragraph(doc, "MADGI — Épargne Santé Retraite", size=17, color=BLUE, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_paragraph(doc, "Comparaison entre le planning cible et les fonctionnalités déjà disponibles",
                  size=12.5, color=GRAY, italic=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    add_callout(
        doc,
        "OBJECTIF.",
        "Servir de boussole commune au Chef de Service ESR, à l’actuaire, à la DSI, "
        "aux équipes financières et au développeur jusqu’à la recette et la mise en production.",
        fill=PALE_BLUE,
        accent=NAVY,
    )
    add_paragraph(doc, "Version 1.0  |  Situation analysée au 25 juillet 2026",
                  size=10, color=GRAY, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=4)
    add_paragraph(doc, "Périmètre : application, documents métier, classeurs actuariels et exigences de réunion",
                  size=9.5, color=GRAY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    doc.add_page_break()

    add_heading(doc, "1. Mode d’emploi de la boussole", 1)
    add_paragraph(
        doc,
        "Ce document distingue ce qui existe réellement dans l’application, ce qui est seulement "
        "documenté, ce qui reste à développer et ce qui ne peut pas avancer sans décision formelle.",
    )
    for label, meaning in [
        ("DISPONIBLE", "fonction utilisable ou document de référence exploitable."),
        ("PARTIEL", "socle existant, mais contrôle, workflow ou calcul incomplet."),
        ("ABSENT", "fonction ou livrable à construire."),
        ("ARBITRAGE", "décision obligatoire avant implémentation définitive."),
    ]:
        color, fill = STATUS[label.title() if label != "ARBITRAGE" else "Arbitrage"]
        table = doc.add_table(rows=1, cols=2)
        set_table_geometry(table, [1.3, 5.2])
        set_cell_shading(table.cell(0, 0), fill)
        set_cell_text(table.cell(0, 0), label, bold=True, color=color, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(0, 1), meaning, size=9.5)
    add_paragraph(doc, "", after=1)

    add_heading(doc, "2. Résumé exécutif", 1)
    add_callout(
        doc,
        "CONSTAT.",
        "Le projet possède un socle applicatif solide : adhérents, bénéficiaires, cotisations, "
        "précomptes, paiements, prestations, rentes, comptes ESR, paramètres, utilisateurs et audit. "
        "La finalisation dépend surtout du gel des règles financières, du moteur actuariel complet, "
        "des workflows de validation, des échanges APS/DGI et des documents officiels.",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )
    summary_rows = [
        ["Arbitrage et cadrage", "50 %", "Sources disponibles, contradictions non tranchées."],
        ["Moteur actuariel", "40 %", "Cotisation avancée ; PM, rachat, décès et rente incomplets."],
        ["Adhérents et bénéficiaires", "75 %", "Fonctions principales présentes ; documents à finaliser."],
        ["Précomptes", "60 %", "Génération présente ; formats officiels et retour DGI manquent."],
        ["Paiements", "45 %", "Saisie présente ; reçus et validations manquent."],
        ["Prestations et rentes", "35 %", "Écrans présents ; workflows et calculs automatiques manquent."],
        ["Reporting", "30 %", "Dashboard présent ; états réglementaires à construire."],
        ["Sécurité et audit", "70 %", "Bonne base ; homologation et exploitation à compléter."],
        ["Migration", "15 %", "Sources reçues ; processus d’import non finalisé."],
        ["Documentation/déploiement", "30 %", "Configuration partielle ; PRA et guides manquent."],
    ]
    add_table(doc, ["Lot", "Estimation", "Lecture"], summary_rows, [2.2, 1.0, 3.3], font_size=9)
    add_paragraph(
        doc,
        "Durée cible : 6 semaines avec disponibilité régulière des décideurs et de l’actuaire ; "
        "7 à 8 semaines avec un seul développeur ou des validations différées.",
        size=10.5, bold=True, color=NAVY,
    )

    doc.add_page_break()
    add_heading(doc, "3. Décisions bloquantes à prendre", 1)
    add_paragraph(
        doc,
        "Ces décisions constituent la porte d’entrée du planning. Aucun développement financier "
        "définitif ne doit être validé avant leur signature.",
    )
    add_table(
        doc,
        ["Sujet", "Décision attendue", "Responsable", "Échéance", "État"],
        arbitrations,
        [1.05, 2.65, 1.15, 0.65, 1.0],
        font_size=8.2,
    )

    add_heading(doc, "4. Comparaison détaillée planning / existant", 1)
    add_paragraph(
        doc,
        "La matrice ci-dessous est la référence de suivi. Un élément PARTIEL ne peut passer à "
        "DISPONIBLE que lorsque son critère de validation métier, technique et documentaire est rempli.",
    )
    for title, rows in comparison_sections:
        add_heading(doc, title, 2)
        add_table(
            doc,
            ["Domaine", "Déjà disponible", "Reste à finaliser", "Statut", "Priorité"],
            rows,
            [1.22, 1.65, 2.18, 0.68, 0.77],
            font_size=7.6,
            status_col=3,
            priority_col=4,
        )

    doc.add_page_break()
    add_heading(doc, "5. Feuille de route de finalisation", 1)
    add_table(
        doc,
        ["Phase", "Lot", "Période", "Pilote", "Critère de sortie"],
        roadmap,
        [0.48, 1.38, 0.92, 1.22, 2.5],
        font_size=8.4,
    )

    add_heading(doc, "5.1 Semaine 1 — Arbitrage et gel", 2)
    for item in [
        "Signer les règles de rachat, décès, invalidité, rente, dates de valeur et arrondis.",
        "Valider l’architecture technique et le périmètre du module Placement.",
        "Recevoir les formats APS/DGI et les données de test anonymisées.",
        "Créer le référentiel versionné des paramètres actuariels.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.2 Semaines 1–2 — Moteur actuariel", 2)
    for item in [
        "Finaliser cotisation unique, PM individuelle/globale, rachat, décès, invalidité et rente.",
        "Supprimer les placeholders et centraliser les règles dans un moteur auditable.",
        "Créer les cas tests de non-régression et faire signer les résultats par l’actuaire.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.3 Semaine 3 — Opérations", 2)
    for item in [
        "Finaliser les précomptes aller/retour APS-DGI et les régularisations.",
        "Ajouter reçu, workflow chèque et mise à jour contrôlée du compte ESR.",
        "Mettre en place les dossiers prestations de bout en bout.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.4 Semaine 4 — Documents et reporting", 2)
    for item in [
        "Produire simulation, bulletin, reçus, liquidations, avis annuel et bordereaux APS.",
        "Générer CIMA C-20, provisions, rentes, capitaux décès et bilan cantonné.",
        "Valider les exports PDF/Excel et leur archivage.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.5 Semaine 5 — Données et sécurité", 2)
    for item in [
        "Importer l’historique, contrôler les doublons et réconcilier les totaux.",
        "Versionner le schéma, tester les droits, sécuriser les secrets et configurer HTTPS.",
        "Tester sauvegarde, restauration, supervision et procédure de reprise.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.6 Semaine 6 — Recette et production", 2)
    for item in [
        "Exécuter la recette métier sur des cas représentatifs.",
        "Corriger les anomalies bloquantes et produire le PV de recette.",
        "Former les profils, mettre en production et activer le suivi renforcé.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Documents et informations encore attendus", 1)
    for item in missing_docs:
        checklist_paragraph = add_bullet(doc, f"☐ {item}")
        checklist_paragraph.paragraph_format.space_after = Pt(1)

    add_heading(doc, "7. Répartition des responsabilités (RACI)", 1)
    add_paragraph(doc, "R = Réalise  |  A = Approuve  |  C = Consulté  |  I = Informé", size=9.5, color=GRAY)
    add_table(
        doc,
        ["Activité", "R", "A", "C", "I"],
        raci,
        [1.55, 1.25, 1.15, 1.55, 1.0],
        font_size=8.3,
    )

    doc.add_page_break()
    add_heading(doc, "8. Registre des risques prioritaires", 1)
    add_table(
        doc,
        ["Risque", "Impact", "Probabilité", "Réponse", "Propriétaire"],
        risks,
        [1.3, 0.82, 0.82, 2.35, 1.21],
        font_size=8.1,
    )

    add_heading(doc, "9. Tableau hebdomadaire de pilotage", 1)
    add_paragraph(
        doc,
        "À renseigner lors de chaque réunion de suivi. Une action n’est terminée que si la preuve "
        "associée est disponible : PV, capture, test, document validé ou rapport.",
    )
    tracking_rows = [
        [str(i), "", "", "", "", "À lancer", ""]
        for i in range(1, 13)
    ]
    add_table(
        doc,
        ["N°", "Action", "Responsable", "Échéance", "Preuve attendue", "Statut", "Blocage"],
        tracking_rows,
        [0.35, 1.55, 1.0, 0.72, 1.18, 0.75, 0.95],
        font_size=7.8,
    )

    doc.add_page_break()
    add_heading(doc, "10. Critères de mise en production", 1)
    acceptance = [
        "Toutes les décisions bloquantes sont signées.",
        "Aucun calcul actuariel ne dépend d’un montant saisi manuellement sans contrôle.",
        "Les tests croisés respectent la tolérance maximale de 0,1 %.",
        "Les flux APS/DGI sont testés sur un aller-retour complet.",
        "Les paiements exigent les validations définies et alimentent correctement le compte ESR.",
        "Les prestations respectent les pièces, calculs, validations et délais.",
        "Les états réglementaires sont réconciliés avec la base et la comptabilité.",
        "Les rôles et droits ont été testés pour chaque profil.",
        "La sauvegarde et la restauration ont été testées avec succès.",
        "Le PV de recette, les guides et le plan de retour arrière sont validés.",
    ]
    for item in acceptance:
        add_bullet(doc, f"☐ {item}")

    add_callout(
        doc,
        "RÈGLE DE PILOTAGE.",
        "Le logiciel ne doit pas être déclaré finalisé parce que les écrans existent. "
        "La finalisation est atteinte lorsque les règles sont signées, les calculs réconciliés, "
        "les workflows contrôlés, les documents produits et la reprise après incident testée.",
        fill=LIGHT_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "11. Sources analysées", 1)
    sources = [
        "Compte rendu de réunion de présentation ESR — juin 2026.",
        "Note technique MADGI Épargne Santé Retraite.",
        "Module Cotisation Épargne Santé Retraite MADGI 2.xlsm.",
        "PM MADGI-ESR exemple.xlsx.",
        "Rachat total ESR.xlsx.",
        "Capital décès agent en activité.xlsx.",
        "Explication du remboursement du capital décès.docx.",
        "Rapport des travaux de l’Équipe Projet — version définitive.",
        "Dépôt applicatif MADGI ESR : frontend React et API Express/TypeScript.",
    ]
    for item in sources:
        add_bullet(doc, item)

    add_paragraph(
        doc,
        "Fin du document — prochaine mise à jour recommandée après la réunion d’arbitrage de la semaine 1.",
        size=9.5, color=GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=18,
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
