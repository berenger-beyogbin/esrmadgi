from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


path = Path(sys.argv[1])
doc = Document(path)

result: dict[str, object] = {
    "core_properties": {
        "title": doc.core_properties.title,
        "subject": doc.core_properties.subject,
        "author": doc.core_properties.author,
        "created": str(doc.core_properties.created),
        "modified": str(doc.core_properties.modified),
    },
    "paragraphs": [],
    "tables": [],
    "headers": [],
    "footers": [],
}

for i, paragraph in enumerate(doc.paragraphs, 1):
    text = paragraph.text.strip()
    if text:
        result["paragraphs"].append(
            {"index": i, "style": paragraph.style.name, "text": text}
        )

for table_index, table in enumerate(doc.tables, 1):
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    result["tables"].append({"index": table_index, "rows": rows})

for section_index, section in enumerate(doc.sections, 1):
    header_text = "\n".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
    footer_text = "\n".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
    if header_text:
        result["headers"].append({"section": section_index, "text": header_text})
    if footer_text:
        result["footers"].append({"section": section_index, "text": footer_text})

with zipfile.ZipFile(path) as archive:
    names = set(archive.namelist())
    result["package_flags"] = {
        "comments": "word/comments.xml" in names,
        "footnotes": "word/footnotes.xml" in names,
        "endnotes": "word/endnotes.xml" in names,
    }
    document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    result["tracked_changes"] = {
        "insertions": len(root.xpath(".//w:ins", namespaces=ns)),
        "deletions": len(root.xpath(".//w:del", namespaces=ns)),
    }

print(json.dumps(result, ensure_ascii=False, indent=2))
