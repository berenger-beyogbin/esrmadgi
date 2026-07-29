from __future__ import annotations

import json
import sys
from pathlib import Path

import pdfplumber
from docx import Document
from openpyxl import load_workbook


root = Path(sys.argv[1])
report: dict[str, object] = {"pdfs": [], "docx": [], "workbooks": []}

for path in sorted(root.iterdir()):
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        pages = []
        with pdfplumber.open(path) as pdf:
            for number, page in enumerate(pdf.pages, 1):
                pages.append({"page": number, "text": page.extract_text() or ""})
        report["pdfs"].append({"file": path.name, "pages": pages})
    elif suffix == ".docx":
        doc = Document(path)
        paragraphs = [
            {"style": p.style.name, "text": p.text.strip()}
            for p in doc.paragraphs
            if p.text.strip()
        ]
        tables = []
        for table in doc.tables:
            tables.append([[cell.text.strip() for cell in row.cells] for row in table.rows])
        report["docx"].append(
            {"file": path.name, "paragraphs": paragraphs, "tables": tables}
        )
    elif suffix in {".xlsx", ".xlsm"}:
        wb_formula = load_workbook(
            path, data_only=False, read_only=False, keep_vba=suffix == ".xlsm"
        )
        wb_values = load_workbook(
            path, data_only=True, read_only=False, keep_vba=suffix == ".xlsm"
        )
        sheets = []
        for ws in wb_formula.worksheets:
            value_ws = wb_values[ws.title]
            cells = []
            formulas = []
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    item = {
                        "cell": cell.coordinate,
                        "value": cell.value,
                        "cached": value_ws[cell.coordinate].value,
                        "number_format": cell.number_format,
                    }
                    cells.append(item)
                    if isinstance(cell.value, str) and cell.value.startswith("="):
                        formulas.append(item)
            sheets.append(
                {
                    "title": ws.title,
                    "state": ws.sheet_state,
                    "max_row": ws.max_row,
                    "max_column": ws.max_column,
                    "merged_ranges": [str(rng) for rng in ws.merged_cells.ranges],
                    "cells": cells,
                    "formulas": formulas,
                }
            )
        report["workbooks"].append(
            {
                "file": path.name,
                "has_vba": bool(getattr(wb_formula, "vba_archive", None)),
                "defined_names": list(wb_formula.defined_names),
                "sheets": sheets,
            }
        )

print(json.dumps(report, ensure_ascii=False, indent=2, default=str))
