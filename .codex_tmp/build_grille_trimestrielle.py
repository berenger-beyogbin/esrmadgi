import csv, math, os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\PROJETS\madgi-esr")
OUT = ROOT / "output" / "Grille_de_cotisation_actuarielle_par_trimestre_MADGI_ESR.docx"
MORTALITE = Path(os.environ["TEMP"]) / "madgi_mortalite.csv"
LOGO = ROOT / "src" / "assets" / "logos" / "logo-madgi.jpg"

BLUE = "2B529F"; DARK = "173B75"; LIGHT = "E8EFFA"; PALE = "F5F8FC"
GOLD = "DF9F28"; WHITE = "FFFFFF"; GRID = "AEBED4"; TEXT = "172033"; MUTED = "5F6B7A"
FONT = "Arial"

with MORTALITE.open(encoding="utf-8-sig") as f:
    LX = {int(r["age_mort"]): float(r["lx"]) for r in csv.DictReader(f)}

def amount(age_retraite, n):
    annual = 600000; rate = 0.035; fees = 0.05; age_max = 106
    ly = LX[age_retraite]
    v = 1 / (1 + rate)
    annuity = sum(LX.get(age_retraite + k, 0) * (v ** k) for k in range(age_max - age_retraite)) / ly
    capital = annual * (1 + fees) * annuity
    iq = (1 + rate) ** 0.25 - 1
    raw = capital * (iq / ((1 + iq) * ((1 + iq) ** n - 1)))
    return int(math.ceil(raw / 100) * 100)

def set_font(run, size=9, bold=False, color=TEXT, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = tc_pr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell, top=30, start=90, bottom=30, end=90):
    tc_pr = cell._tc.get_or_add_tcPr(); tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None: tc_mar = OxmlElement("w:tcMar"); tc_pr.append(tc_mar)
    for tag, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None: node = OxmlElement(f"w:{tag}"); tc_mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")

def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr(); tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None: tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa)); tc_w.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW")); tbl_w.set(qn("w:w"), str(sum(widths))); tbl_w.set(qn("w:type"), "dxa")
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None: ind = OxmlElement("w:tblInd"); tbl_pr.append(ind)
    ind.set(qn("w:w"), "90"); ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells): set_cell_width(cell, widths[i]); cell_margins(cell)

def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr(); hdr = OxmlElement("w:tblHeader"); hdr.set(qn("w:val"), "true"); tr_pr.append(hdr)

def add_page_field(paragraph):
    run = paragraph.add_run(); begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"; end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end): run._r.append(node)
    set_font(run, 8, color=MUTED)

def duration(n):
    years, quarters = divmod(n, 4)
    if quarters == 0: return f"{years} an" + ("s" if years > 1 else "")
    y = f"{years} an" + ("s" if years > 1 else "") if years else "0 an"
    return f"{y} et {quarters} trimestre" + ("s" if quarters > 1 else "")

doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width = Inches(11.6929); sec.page_height = Inches(8.2677)
sec.top_margin = Inches(0.38); sec.bottom_margin = Inches(0.42)
sec.left_margin = Inches(0.55); sec.right_margin = Inches(0.55)
sec.header_distance = Inches(0.18); sec.footer_distance = Inches(0.18)

normal = doc.styles["Normal"]
normal.font.name = FONT; normal._element.rPr.rFonts.set(qn("w:ascii"), FONT); normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(9); normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_after = Pt(0); normal.paragraph_format.line_spacing = 1.0

header = sec.header
hp = header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(hp.add_run("ESR-MADGI  |  Grille actuarielle trimestrielle"), 8, True, BLUE)
footer = sec.footer
fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(fp.add_run("Hypothèses : rente annuelle 600 000 FCFA • couverture 100 % • taux technique 3,5 % • frais de rente 5 % • table CIMA-F • âge maximal 106 ans • arrondi supérieur à 100 FCFA    |    Page "), 7.5, color=MUTED)
add_page_field(fp)

widths = [1800, 3450, 5000, 5000]
for page_index, start in enumerate((1, 31, 61, 91)):
    if page_index:
        boundary = doc.add_paragraph()
        boundary.paragraph_format.page_break_before = True
        boundary.paragraph_format.space_before = Pt(0)
        boundary.paragraph_format.space_after = Pt(0)
        boundary.paragraph_format.line_spacing = Pt(1)
    title = doc.add_table(rows=1, cols=2)
    title.style = "Table Grid"; title.autofit = False
    title.cell(0,0).width = Inches(0.8); title.cell(0,1).width = Inches(9.65)
    title.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists(): title.cell(0,0).paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(0.42))
    shade(title.cell(0,0), WHITE); shade(title.cell(0,1), WHITE)
    p = title.cell(0,1).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    set_font(p.add_run("GRILLE DE COTISATION ACTUARIELLE PAR TRIMESTRE"), 13.5, True, BLUE)
    p2 = title.cell(0,1).add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p2.add_run(f"Lecture directe — trimestres {start} à {start+29}"), 8.5, True, GOLD)
    for row in title.rows:
        for cell in row.cells:
            tc_pr=cell._tc.get_or_add_tcPr(); borders=tc_pr.find(qn("w:tcBorders"))
            if borders is None: borders=OxmlElement("w:tcBorders"); tc_pr.append(borders)
            for edge in ("top","left","bottom","right","insideH","insideV"):
                el=OxmlElement(f"w:{edge}"); el.set(qn("w:val"),"nil"); borders.append(el)

    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(0); note.paragraph_format.space_after = Pt(1)
    set_font(note.add_run("Mode d'emploi : repérer le nombre exact de trimestres jusqu'à la retraite, puis lire la cotisation correspondant à l'âge de départ. Les lignes bleues signalent les années complètes."), 7.5, color=DARK)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["TRIMESTRES\nRESTANTS", "DURÉE ÉQUIVALENTE", "COTISATION TRIMESTRIELLE\nRETRAITE À 60 ANS", "COTISATION TRIMESTRIELLE\nRETRAITE À 65 ANS"]
    for i, text in enumerate(headers):
        cell=table.rows[0].cells[i]; shade(cell, BLUE); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(text), 8.5, True, WHITE)
    repeat_header(table.rows[0])
    table.rows[0].height = Inches(0.32)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for n in range(start, start + 30):
        row = table.add_row(); annual = n % 4 == 0
        row.height = Inches(0.19)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        values = [str(n), duration(n), f"{amount(60,n):,}".replace(","," ") + " FCFA", f"{amount(65,n):,}".replace(","," ") + " FCFA"]
        for i, value in enumerate(values):
            cell=row.cells[i]; cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cell, LIGHT if annual else (PALE if n % 2 == 0 else WHITE))
            p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.RIGHT
            set_font(p.add_run(value), 8.3, annual, DARK if annual else TEXT)
    set_table_geometry(table, widths)

doc.core_properties.title = "Grille de cotisation actuarielle par trimestre MADGI ESR"
doc.core_properties.subject = "Cotisations trimestrielles pour 1 à 120 trimestres avant la retraite"
doc.core_properties.author = "MADGI ESR"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
print("reference checks", amount(60,4), amount(60,40), amount(60,120), amount(65,4), amount(65,40), amount(65,120))
